import os
import re
from docopt import docopt
from bs4 import BeautifulSoup
from docx import Document
from langchain.schema import Document as LangChainDocument
import ollama
import markdown
from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_ollama import ChatOllama
from langchain import hub
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains.retrieval import create_retrieval_chain 
import time
from rag_tester import load_json 

llm =ChatOllama(model="gemma:7b")
embed_model="ordis/jina-embeddings-v2-base-code:latest"

"""
pip install -qU langchain_community faiss-cpu

"""
def extract_text_from_html(file_path):
    """
    Extracts text content from an HTML file, removing script and style elements.

    Args:
        file_path (str): Path to the HTML file.

    Returns:
        str: Cleaned text extracted from the HTML file.
    """
    with open(file_path, "r", encoding="utf-8") as file:
        soup = BeautifulSoup(file, "html.parser")
        for script_or_style in soup(["script", "style"]):
            script_or_style.decompose()
        return soup.get_text()


def extract_text_from_docx(file_path):
    """
    Extracts text from a DOCX file.

    Args:
        file_path (str): Path to the DOCX file.

    Returns:
        str: Text content of the DOCX file.
    """
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text_from_markdown(file_path):
    """
    Extracts text from a Markdown file.

    Args:
        file_path (str): Path to the Markdown file.

    Returns:
        str: Text content of the Markdown file.
    """
    with open(file_path, "r", encoding="utf-8") as file:
        html = markdown.markdown(file.read())
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text()


def extract_text_from_txt(file_path):
    """
    Extracts text from a plain text file.

    Args:
        file_path (str): Path to the text file.

    Returns:
        str: Text content of the text file.
    """
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()


def process_files(directory_path):
    """
    Processes all supported file types in a specified directory by extracting text from each.

    Args:
        directory_path (str): Path to the directory containing files.

    Returns:
        list: A list of dictionaries, each containing the text and metadata of a file.
    """
    documents = []
    for file_name in os.listdir(directory_path):
        file_path = os.path.join(directory_path, file_name)

        match file_name.split('.')[-1]:
            case "html":
                text = extract_text_from_html(file_path)
            case "docx":
                text = extract_text_from_docx(file_path)
            case "md":
                text = extract_text_from_markdown(file_path)
            case "txt":
                text = extract_text_from_txt(file_path)
            case _:
                print(f"Unsupported file type: {file_name}")
                continue
        documents.append( LangChainDocument(page_content=text,metadata={"source": file_name}))
        #documents.append({"text": text, "metadata": {"source": file_name}})

    return documents

def split_text_into_chunks(documents, chunk_size, chunk_overlap):
    """
    Splits the text content of documents into smaller chunks.

    Args:
        documents (list): A list of dictionaries containing text and metadata.
        chunk_size (int): Maximum size of each text chunk.
        chunk_overlap (int): Overlap between consecutive chunks.

    Returns:
        list: A list of dictionaries, each containing a chunk of text and associated metadata.
    """
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=int(chunk_size), chunk_overlap=int(chunk_overlap))

    return text_splitter.split_documents(documents)


def create_vector_store(processed_documents):
    embeddings = OllamaEmbeddings(model=embed_model)
    return FAISS.from_documents(processed_documents, embeddings)


def create_rag_system(vector_store):
    """
    Creates a Retrieval-Augmented Generation (RAG) system by integrating a retriever and a local LLM served by Ollama.

    Args:
        vector_store (FAISS): The FAISS vector store to use as the retriever.
        model_name (str): Name of the language model to use for text generation (e.g., "llama2").

    Returns:
        RetrievalQA: A RAG system capable of answering queries using the vector store and Ollama local language model.
    """
    # Step 1: Set up the retriever (FAISS vector store)
    retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 3})
    retrieval_qa_chat_prompt = hub.pull("langchain-ai/retrieval-qa-chat")
    combine_docs_chain = create_stuff_documents_chain(llm, retrieval_qa_chat_prompt)
    rag_chain = create_retrieval_chain(retriever, combine_docs_chain)
    return rag_chain

import re


if __name__ == "__main__":
    """
    Main script to process files, set up the RAG system, and answer queries interactively.
    """
    # Parse command-line arguments
    chunk_size=400
    
    """
     Smaller chunks are better for detailed analysis, 
     while larger chunks are suitable for broader understanding,
    
    """
    chunk_overlap=100
    
    pwd = os.getcwd()
    input_dir=pwd + "/paypal_rag/rag-data"

    # Step 1-2: Extract and process files
    documents = process_files(input_dir)

    # Step 3: Split text into chunks
    processed_documents = split_text_into_chunks(documents, chunk_size, chunk_overlap)

    # Step 4: Create vector store
    vector_store = create_vector_store(processed_documents)

    # Step 5: Create RAG system
    qa_chain = create_rag_system(vector_store)

    # # Step 6: Query the RAG system
    print("RAG System is ready! Type your query below (or type 'exit' to quit):")
    while True:
        query = input("Query: ")
        if query.lower() == "exit":
            print("Exiting...")
            break
        result = qa_chain.invoke({"input": query})
        print("Answer:", result['answer'])

    # Step 7: Validate the model performance
    print("Running Tests!")
    # Example usage
    file_path = pwd + "/paypal_rag/rag-data/ladakh-questions.json"
    qa_data = load_json(file_path)
    if qa_data:
        for key, value in qa_data.items():
            print(f"Q{key}: {value['question']}")
            print(f"Expected{key}: {value['answer']}\n")
            result = qa_chain.invoke({"input": value['question']})
            print(f"Actual: {result['answer']}\n")   
